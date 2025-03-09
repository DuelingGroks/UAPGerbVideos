import os
import markdown
from docx import Document

def convert_markdown_to_word(obsidian_folder, output_docx):
    if not os.path.exists(obsidian_folder):
        os.makedirs(obsidian_folder)
    
    doc = Document()
    doc.add_heading('Obsidian Notes Compilation', level=1)
    
    for root, _, files in os.walk(obsidian_folder):
        for filename in files:
            if filename.endswith(".md"):
                filepath = os.path.join(root, filename)
                print(f"Processing file: {filepath}")  # Debugging
                try:
                    with open(filepath, 'r', encoding='utf-8') as file:
                        md_content = file.read()
                        if not md_content.strip():
                            print(f"Warning: {filename} is empty.")
                            continue
                        
                        relative_path = os.path.relpath(filepath, obsidian_folder)
                        doc.add_heading(relative_path.replace('.md', ''), level=2)
                        doc.add_paragraph(md_content)
                        doc.add_page_break()
                except Exception as e:
                    print(f"Error reading {filename}: {e}")
    
    doc.save(output_docx)
    print(f"Conversion complete: {output_docx}")

def generate_directory_listing(obsidian_folder, output_txt):
    if not os.path.exists(obsidian_folder):
        os.makedirs(obsidian_folder)
    
    with open(output_txt, 'w', encoding='utf-8') as file:
        file.write("Directory Listing for Obsidian Folder:\n\n")
        for root, dirs, files in os.walk(obsidian_folder):
            file.write(f"{root}:\n")
            for filename in files:
                file.write(f"  - {filename}\n")
    print(f"Directory listing saved: {output_txt}")

if __name__ == "__main__":
    obsidian_folder = os.path.join(os.getcwd(), "UAPGerb")  # Absolute path
    output_docx = "obsidian_notes_compilation.docx"
    output_txt = "directory_listing.txt"
    
    generate_directory_listing(obsidian_folder, output_txt)
    convert_markdown_to_word(obsidian_folder, output_docx)
