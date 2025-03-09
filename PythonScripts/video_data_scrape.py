import subprocess
import json
import os
import re
import sys
from docx import Document


def sanitize_filename(filename):
    """
    Removes or replaces characters not allowed in filenames for most OSes.
    """
    invalid_chars = '<>:"/\\|?*'
    for ch in invalid_chars:
        filename = filename.replace(ch, '')
    filename = filename.strip().rstrip('.')
    return filename


def download_metadata(url):
    """
    Runs yt-dlp in JSON mode to get metadata about the video.
    Returns a dictionary containing metadata (title, description, webpage_url, etc.).
    """
    cmd = ["yt-dlp", "-J", "--skip-download", url]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise Exception(f"yt-dlp failed to get metadata for {url}:\n{result.stderr}")
    info = json.loads(result.stdout)
    return info


def download_thumbnail(url, outtmpl, convert_format="jpg"):
    """
    Downloads the video thumbnail using yt-dlp.
    Saves the thumbnail with a filename based on 'outtmpl'.
    Returns the thumbnail filename if successful, or None otherwise.
    """
    cmd = [
        "yt-dlp",
        "--skip-download",
        "--write-thumbnail",
        "--convert-thumbnails", convert_format,
        "-o", outtmpl,
        url
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"Warning: yt-dlp could not download thumbnail. Stderr:\n{result.stderr}")
        return None
    
    thumbnail_file = outtmpl.replace("%(ext)s", convert_format)
    if os.path.isfile(thumbnail_file):
        return thumbnail_file
    else:
        return None


def download_subtitles(url, output_directory=None):
    """
    Downloads the automatic subtitles (all languages) in .vtt format using yt-dlp.
    We then look for the created 'subtitles.*.vtt' file in the output directory.
    Returns the found subtitle filename if any, or None.
    """
    # Build an output template in the chosen directory
    if output_directory:
        outtmpl = os.path.join(output_directory, "subtitles.%(id)s.%(ext)s")
    else:
        outtmpl = "subtitles.%(id)s.%(ext)s"
    
    cmd = [
        "yt-dlp",
        "--write-auto-subs",
        "--skip-download",
        "--sub-langs", "en",          # or you can do '--sub-langs', 'en.*' if you only want English variants
        "--convert-subs", "vtt",
        "-o", outtmpl,
        url
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"Warning: yt-dlp could not download subtitles. Stderr:\n{result.stderr}")
        return None
    
    # Search for subtitles.*.vtt in the output directory
    search_dir = output_directory if output_directory else "."
    found_vtt = None
    for f in os.listdir(search_dir):
        if f.startswith("subtitles.") and f.endswith(".vtt"):
            candidate_path = os.path.join(search_dir, f)
            if os.path.isfile(candidate_path):
                found_vtt = candidate_path
                break

    return found_vtt


def parse_vtt_subtitles(vtt_filename):
    """
    Parses a .vtt file to extract transcript text.
    Removes timestamps, <c> tags, and lines with -->.
    Returns a cleaned transcript as a single string.
    """
    transcript_lines = []
    with open(vtt_filename, "r", encoding="utf-8") as f:
        for line in f:
            # Remove HTML/markup tags like <c>, <00:00:00.000>, etc.
            line = re.sub(r"<[^>]+>", "", line)
            line_stripped = line.strip()
            # Skip empty lines or lines with --> (timing info).
            if not line_stripped or "-->" in line_stripped:
                continue
            transcript_lines.append(line_stripped)
    # Join into one big paragraph and normalize spacing
    transcript_text = " ".join(transcript_lines)
    transcript_text = re.sub(r"\s+", " ", transcript_text).strip()
    return transcript_text


def create_word_document(title, description, link, transcript, output_docx="output.docx"):
    """
    Creates a Word document with the provided video info and transcript.
    Sets the document's internal title to 'title'.
    """
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)
    
    document.add_heading("Description", level=1)
    document.add_paragraph(description)

    document.add_heading("Video Link", level=1)
    document.add_paragraph(link)
    
    document.add_heading("Transcript", level=1)
    document.add_paragraph(transcript)
    
    document.save(output_docx)


def download_video(url, output_directory=None, safe_title=None):
    """
    Downloads the video itself using yt-dlp if requested with -v.
    """
    # If we have a known safe_title, we can use it for naming:
    #   e.g. "MyVideoTitle.%(ext)s" inside the chosen directory.
    # Or you can let yt-dlp use its own naming scheme. Adjust as you prefer.
    if not safe_title:
        outtmpl = os.path.join(output_directory, "%(title).%(ext)s") if output_directory else "%(title).%(ext)s"
    else:
        outtmpl = os.path.join(output_directory, safe_title + ".%(ext)s") if output_directory else safe_title + ".%(ext)s"

    cmd = [
        "yt-dlp",
        "-f", "best",          # or whichever format you want
        "-o", outtmpl,
        url
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Warning: Failed to download video. Stderr:\n{result.stderr}")


def process_video(url, output_directory=None, download_vid=False):
    print(f"\nProcessing URL: {url}")
    
    # 1. Retrieve metadata
    info = download_metadata(url)
    title = info.get("title", "No Title Found").strip()
    description = info.get("description", "No Description Found").strip()
    webpage_url = info.get("webpage_url", url).strip()
    
    # 2. Generate safe filename base from title
    safe_title = sanitize_filename(title)
    if not safe_title:
        safe_title = "video_info"  # fallback if the title becomes empty
    
    # 3. (Optional) Download the video
    if download_vid:
        download_video(url, output_directory=output_directory, safe_title=safe_title)
    
    # 4. Download subtitles and parse transcript (if available)
    vtt_file = download_subtitles(url, output_directory=output_directory)
    if vtt_file and os.path.isfile(vtt_file):
        transcript = parse_vtt_subtitles(vtt_file)
        print("Transcript successfully parsed.")
    else:
        transcript = "No transcript available or subtitles not found."
    
    # 5. Download the thumbnail
    if output_directory:
        thumbnail_outtmpl = os.path.join(output_directory, safe_title + ".%(ext)s")
    else:
        thumbnail_outtmpl = safe_title + ".%(ext)s"
    
    thumbnail_file = download_thumbnail(url, thumbnail_outtmpl)
    if thumbnail_file:
        print(f"Thumbnail saved as '{thumbnail_file}'.")
    else:
        print("Thumbnail could not be downloaded.")
    
    # 6. Create the Word document
    output_docx = os.path.join(output_directory, f"{safe_title}.docx") if output_directory else f"{safe_title}.docx"
    create_word_document(title, description, webpage_url, transcript, output_docx=output_docx)
    print(f"Word document '{output_docx}' created successfully.")
    
    # 7. Clean up the subtitles file (if desired)
    #    Comment out if you want to keep the .vtt
    if vtt_file and os.path.isfile(vtt_file):
        os.remove(vtt_file)


def print_usage():
    usage_text = """
Usage:
  python script.py [YouTube_URL] [options]

Examples:
  1) Single video by URL:
     python script.py "https://www.youtube.com/watch?v=XYZ" -v -d outdir
       -v  -> download the video
       -d  -> specify output directory

  2) Multiple videos from file:
     python script.py -file urls.txt -v

  3) If no URL is given, -file must be provided:
     python script.py -file urls.txt -d outdir

Options:
  -file FILENAME   Read one URL per line from the specified file.
  -v               Download the video itself (in addition to subtitles, metadata, etc.).
  -d DIRECTORY     Create/use DIRECTORY for all output files.
"""
    print(usage_text)


def main():
    # Basic argument parsing
    # We'll gather:
    #   -- a single URL if present and not prefixed by '-'
    #   -- optional '-file <filename>'
    #   -- optional '-v'
    #   -- optional '-d <directory>'
    args = sys.argv[1:]
    if not args:
        print_usage()
        return

    url = None
    url_list = []
    file_arg = None
    download_vid = False
    output_directory = None

    i = 0
    while i < len(args):
        arg = args[i]
        if arg.startswith("-"):
            # Parse known flags
            if arg == "-file":
                i += 1
                if i < len(args):
                    file_arg = args[i]
                else:
                    print("Error: '-file' requires a filename.")
                    return
            elif arg == "-v":
                download_vid = True
            elif arg == "-d":
                i += 1
                if i < len(args):
                    output_directory = args[i]
                else:
                    print("Error: '-d' requires a directory name.")
                    return
            else:
                print(f"Unknown option '{arg}'")
                print_usage()
                return
        else:
            # Assume it's a URL if not prefixed with '-'
            url = arg
        i += 1

    # If user specified an output directory, create it if not exists
    if output_directory:
        os.makedirs(output_directory, exist_ok=True)

    # If a direct URL was provided, add it to the list
    if url is not None:
        url_list.append(url)

    # If a file was provided, read from that file
    if file_arg:
        try:
            with open(file_arg, "r", encoding="utf-8") as f:
                lines = [line.strip() for line in f if line.strip()]
            url_list.extend(lines)
        except Exception as e:
            print(f"Error reading file '{file_arg}': {e}")
            return

    # If no URL and no file, print usage and exit
    if not url_list:
        print("No URLs provided via command line or file.")
        print_usage()
        return

    # Process each URL
    for each_url in url_list:
        try:
            process_video(each_url, output_directory=output_directory, download_vid=download_vid)
        except Exception as e:
            print(f"Error processing URL '{each_url}': {e}")


if __name__ == "__main__":
    main()
