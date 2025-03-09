import os
import markdown
from docx import Document

def convert_markdown_to_word(obsidian_folder, output_docx):
    doc = Document()
    doc.add_heading('Obsidian Notes Compilation', level=1)
    
    for root, _, files in os.walk(obsidian_folder):
        for filename in files:
            if filename.endswith(".md"):
                filepath = os.path.join(root, filename)
                with open(filepath, 'r', encoding='utf-8') as file:
                    md_content = file.read()
                    doc.add_heading(filepath.replace('.md', ''), level=2)
                    doc.add_paragraph(md_content)
                    doc.add_page_break()
    
    doc.save(output_docx)

def generate_directory_listing(obsidian_folder, output_txt):
    with open(output_txt, 'w', encoding='utf-8') as file:
        for root, _, files in os.walk(obsidian_folder):
            file.write(f"{root}:\n")
            for filename in files:
                file.write(f"  - {filename}\n")

if __name__ == "__main__":
    obsidian_folder = os.path.join(os.getcwd(), "UAPGerb")
    generate_directory_listing(obsidian_folder, "directory_listing.txt")
    convert_markdown_to_word(obsidian_folder, "obsidian_notes_compilation.docx")
